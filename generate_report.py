cat << 'EOF' > generate_report.py

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_image_placeholder(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F0F4F8")
    
    # Border styling
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="6" w:space="0" w:color="1A5276"/>
            <w:left w:val="single" w:sz="6" w:space="0" w:color="1A5276"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="1A5276"/>
            <w:right w:val="single" w:sz="6" w:space="0" w:color="1A5276"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"📷 [ EMPLACEMENT CAPTURE D'ÉCRAN ]\n{text}")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    run.font.size = Pt(10)
    doc.add_paragraph() # Spacing

doc = docx.Document()

# Styles
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run("RAPPORT DE PROJET BIG DATA & NOSQL\n")
r_title.font.size = Pt(20)
r_title.font.bold = True
r_title.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

r_sub = p_title.add_run("Analyse et Visualisation des Retards de Vols avec PySpark & Apache Cassandra")
r_sub.font.size = Pt(13)
r_sub.font.italic = True
r_sub.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

doc.add_paragraph()

# Meta Table
table_meta = doc.add_table(rows=2, cols=2)
table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    [("Niveau :", " Licence 3 Data Science & Big Data"), ("Module :", " Cours Big Data")],
    [("Auteur :", " Mr Corazon"), ("Technologies :", " PySpark, Cassandra, Docker, Streamlit")]
]
for row_idx, row in enumerate(meta_data):
    for col_idx, (label, val) in enumerate(row):
        cell = table_meta.cell(row_idx, col_idx)
        set_cell_background(cell, "EBEDEF")
        p = cell.paragraphs[0]
        r1 = p.add_run(label)
        r1.bold = True
        p.add_run(val)

doc.add_paragraph()

# Sections
def add_h1(text):
    h = doc.add_heading(text, level=1)
    h.style.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    h.style.font.size = Pt(15)
    return h

def add_h2(text):
    h = doc.add_heading(text, level=2)
    h.style.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    h.style.font.size = Pt(12)
    return h

# 1. Contexte
add_h1("1. Contexte et Objectifs du Projet")
p = doc.add_paragraph(
    "L'industrie aéronautique génère quotidiennement d'immenses volumes de données d'exploitation. "
    "L'analyse des retards de vols constitue un enjeu stratégique majeur pour l'optimisation des flux aéroportuaires, "
    "la gestion des flottes et la satisfaction des passagers. Dans le cadre de ce projet Big Data, notre objectif est "
    "de concevoir une chaîne de traitement distribuée et un stockage NoSQL performant pour analyser un jeu de données "
    "de plus de 100 000 enregistrements de vols."
)
p = doc.add_paragraph(
    "Les métriques clés calculées permettent de mesurer l'impact des retards selon trois axes fondamentaux :"
)
bp1 = doc.add_paragraph(style='List Bullet')
bp1.add_run("Par aéroport d'origine : ").bold = True
bp1.add_run("Identifier les hubs générant les plus fortes perturbations.")

bp2 = doc.add_paragraph(style='List Bullet')
bp2.add_run("Par compagnie aérienne : ").bold = True
bp2.add_run("Évaluer la régularité et la durée moyenne des retards par transporteur.")

bp3 = doc.add_paragraph(style='List Bullet')
bp3.add_run("Par jour de la semaine : ").bold = True
bp3.add_run("Mettre en évidence la saisonnalité hebdomadaire du trafic.")

# 2. Architecture
add_h1("2. Architecture et Choix Techniques")
p = doc.add_paragraph(
    "Pour répondre aux exigences de scalabilité et de traitement à faible latence, nous avons mis en place "
    "une architecture articulée autour de quatre composantes principales :"
)

table_tech = doc.add_table(rows=5, cols=2)
table_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Composante", "Rôle et Justification Technique"]
for i, h in enumerate(headers):
    cell = table_tech.cell(0, i)
    set_cell_background(cell, "1A5276")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

techs = [
    ("Apache Cassandra", "Base de données NoSQL orientée colonnes offrant une haute disponibilité, une scalabilité horizontale linéaire et des lectures ultra-rapides sur clés de partition."),
    ("Apache Spark (PySpark 3.5)", "Moteur de calcul distribué en mémoire idéal pour exécuter les opérations de nettoyage, de filtrage et d'agrégation sur de volumineux jeux de données."),
    ("Docker / Docker Compose", "Conteneurisation de l'instance Cassandra garantissant un environnement d'exécution isolé, reproductible et rapidement déployable."),
    ("Streamlit & Plotly", "Framework Python de visualisation de données permettant de créer un dashboard interactif interrogeant directement Apache Cassandra.")
]

for row_idx, (tech, desc) in enumerate(techs, start=1):
    c1 = table_tech.cell(row_idx, 0)
    c2 = table_tech.cell(row_idx, 1)
    if row_idx % 2 == 0:
        set_cell_background(c1, "F9EBEA")
        set_cell_background(c2, "F9EBEA")
    c1.paragraphs[0].add_run(tech).bold = True
    c2.paragraphs[0].add_run(desc)

doc.add_paragraph()

# 3. Schéma NoSQL
add_h1("3. Modélisation du Schéma de Données NoSQL (Cassandra)")
p = doc.add_paragraph(
    "Contrairement aux bases de données relationnelles traditionnelles, la modélisation sous Apache Cassandra repose sur le principe "
    "'Query-First' (conception guidée par les requêtes). Chaque table est spécifiquement optimisée pour répondre à un motif d'accès précis sans nécessiter de jointures coûteuses."
)

add_h2("3.1 Structure du Keyspace et des Tables")
p_cql = doc.add_paragraph(
    "Keyspace créé : flight_analytics (Facteur de réplication : 1, SimpleStrategy).\n\n"
    "• delays_by_airport : Clé de partition (airport_code), Clé de clustering (flight_date DESC).\n"
    "• delays_by_airline : Clé primaire (airline).\n"
    "• delays_by_weekday : Clé primaire (day_of_week)."
)

add_image_placeholder(doc, "Insérer ici la capture d'écran de l'exécution de cqlsh montrant les tables créées (DESCRIBE TABLES).")

# 4. Traitement PySpark
add_h1("4. Traitement et Agrégation des Données avec PySpark")
p = doc.add_paragraph(
    "Le traitement des données s'effectue via le script process_flights.py. L'analyse exploratoire a révélé une "
    "spécificité majeure du jeu de données Kaggle : tous les vols enregistrés présentent au moins 15 minutes de retard (ArrDelay >= 15). "
    "Pour produire une métrique d'analyse pertinente, nous avons défini le Taux de Retard Lourd (> 60 minutes)."
)

add_h2("4.1 Métriques Calculées")
p_m = doc.add_paragraph(
    "1. Total des vols : Nombre total d'enregistrements par groupe.\n"
    "2. Temps moyen de retard : Moyenne de la colonne ArrDelay (en minutes).\n"
    "3. Taux de retard lourd : Proportion de vols ayant un retard supérieur à 60 minutes sur le total des vols."
)

add_image_placeholder(doc, "Insérer ici la capture du terminal montrant la confirmation d'écriture dans Cassandra ('Mise à jour réussie...').")

# 5. Résultats et Visualisation
add_h1("5. Résultats Obtenus et Visualisation (Dashboard Streamlit)")
p = doc.add_paragraph(
    "Afin de valoriser les résultats stockés dans Cassandra, une application Streamlit interroge la base NoSQL en temps réel "
    "via le driver cassandra-driver et restitue des graphiques interactifs Plotly."
)

add_image_placeholder(doc, "Insérer ici la capture d'écran du Dashboard Streamlit - Graphique des retards par Compagnie et par Jour.")

add_image_placeholder(doc, "Insérer ici la capture d'écran du Dashboard Streamlit - Graphique d'analyse temporelle par Aéroport.")

# 6. Difficultés Rencontrées
add_h1("6. Difficultés Rencontrées et Solutions Apportées")

t_diff = doc.add_table(rows=4, cols=3)
t_diff.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Problème rencontré", "Cause identifiée", "Solution appliquée"]
for i, h in enumerate(headers):
    cell = t_diff.cell(0, i)
    set_cell_background(cell, "1A5276")
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

diffs = [
    ("Incompatibilité PySpark / Cassandra", "Conflit de version Scala entre PySpark 4.0 et le connecteur spark-cassandra-connector.", "Downgrade vers PySpark 3.5.0 pour assurer la compatibilité ascendante avec Scala 2.12."),
    ("Biais statistique du Dataset", "Les vols ont tous au moins 15 min de retard, faussant le taux de retard classique.", "Introduction du seuil de retard lourd (> 60 min) pour obtenir une métrique représentative."),
    ("Erreur de colonne Cassandra", "Inadéquation entre le DataFrame Spark (delayed_flights) et la table Cassandra.", "Alignement des schémas en retirant l'alias superflus dans PySpark.")
]

for row_idx, (p_item, c_item, s_item) in enumerate(diffs, start=1):
    c1, c2, c3 = t_diff.cell(row_idx, 0), t_diff.cell(row_idx, 1), t_diff.cell(row_idx, 2)
    if row_idx % 2 == 0:
        for c in (c1, c2, c3):
            set_cell_background(c, "F2F4F4")
    c1.paragraphs[0].add_run(p_item).bold = True
    c2.paragraphs[0].add_run(c_item)
    c3.paragraphs[0].add_run(s_item)

doc.add_paragraph()

# 7. Conclusion
add_h1("7. Conclusion et Perspectives")
p = doc.add_paragraph(
    "Ce projet a permis de valider avec succès le déploiement d'une chaîne Big Data complète. "
    "L'association de PySpark pour le traitement distribué et d'Apache Cassandra pour le stockage orienté colonnes "
    "démontre une efficacité remarquable pour l'analyse de gros volumes de données aéronautiques."
)
p = doc.add_paragraph(
    "En perspective, ce système pourrait être étendu vers du traitement en temps réel en intégrant Apache Kafka "
    "et Spark Streaming afin d'ingérer les flux de vols directement au moment de leur atterrissage."
)

doc.save("Rapport_Projet_BigData.docx")
print("✅ Le fichier 'Rapport_Projet_BigData.docx' a été généré avec succès !")
